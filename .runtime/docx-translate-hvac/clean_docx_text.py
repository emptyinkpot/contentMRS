from docx import Document
import os
import re

ROOT = os.path.dirname(__file__)
SRC = os.path.join(ROOT, "source.docx")
OUT = os.path.join(ROOT, "hvac_cn_layout_preserved.docx")

REPLACEMENTS = {
    "NomENcLATuRE": "符号说明",
    "NOMENCLATURE": "符号说明",
    "TABLE I": "表 I",
    "TABLE II": "表 II",
    "TABLE III": "表 III",
    "TABLE IV": "表 IV",
    "B. Disturbance Compensated Q-Function": "B. 扰动补偿 Q 函数",
    "V. ImpLEmEntAtion, REsuLts, And Discussions": "V. 实现、结果与讨论",
    "VI. ConcLusions": "VI. 结论",
    "REFErEncEs": "参考文献",
    "REFERENCES": "参考文献",
    "Algorithm 1": "算法 1",
    "NIST": "NIST",
}


def has_object(paragraph):
    xml = paragraph._p.xml
    return any(marker in xml for marker in [
        "<w:drawing",
        "<wp:inline",
        "<wp:anchor",
        "<m:oMath",
        "<m:oMathPara",
        "<w:object",
        "<v:shape",
        "<w:pict",
    ])


def replace_preserve_first_run(paragraph, new_text):
    if has_object(paragraph):
        return False
    if not paragraph.runs:
        paragraph.add_run(new_text)
        return True
    for run in paragraph.runs:
        run.text = ""
    paragraph.runs[0].text = new_text
    return True


def clean_text(text):
    next_text = text
    for src, dst in REPLACEMENTS.items():
        next_text = next_text.replace(src, dst)
    next_text = re.sub(r"\bDisturbance Compensated Q-Function\b", "扰动补偿 Q 函数", next_text)
    next_text = re.sub(r"\bBellman 方程的求解器\b", "贝尔曼方程的求解器", next_text)
    next_text = next_text.replace("occupants", "居住者")
    return next_text


def clean_paragraphs(paragraphs):
    changed = 0
    for paragraph in paragraphs:
        original = paragraph.text
        if not original:
            continue
        updated = clean_text(original)
        if updated != original:
            if replace_preserve_first_run(paragraph, updated):
                changed += 1
    return changed


def main():
    doc = Document(SRC)
    changed = clean_paragraphs(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                changed += clean_paragraphs(cell.paragraphs)
    doc.save(OUT)
    print(f"saved={OUT}")
    print(f"changed={changed}")


if __name__ == "__main__":
    main()
